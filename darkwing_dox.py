# darkwing_dox.py
# Darkwing Dox v3.3 – Webigail Rising
# 🦇 The caped ripper of the documentverse

import os
import re
import sys
import zipfile
import shutil
import subprocess
import tkinter as tk
from tkinter import filedialog, messagebox
from datetime import datetime
from urllib.parse import urlsplit
from urllib.request import urlretrieve
import requests

from docx import Document
from PIL import Image
import pytesseract

# Paths for external tools
LIBREOFFICE_PATH = r"C:\\Program Files\\LibreOffice\\program\\soffice.exe"
TESSERACT_PATH = r"C:\\Program Files\\Tesseract-OCR\\tesseract.exe"
pytesseract.pytesseract.tesseract_cmd = TESSERACT_PATH

def spoofed_download(url, save_dir):
    local_filename = os.path.basename(urlsplit(url).path)
    local_path = os.path.join(save_dir, local_filename)
    headers = {'Referer': url}  # Webigail spoof mode
    try:
        r = requests.get(url, headers=headers, verify=False, timeout=30)
        r.raise_for_status()
        with open(local_path, 'wb') as f:
            f.write(r.content)
        return local_path
    except Exception as e:
        print(f"[!] Failed to download {url}: {e}")
        return None

def convert_doc_to_docx(doc_path, output_dir):
    subprocess.run([LIBREOFFICE_PATH, "--headless", "--convert-to", "docx", "--outdir", output_dir, doc_path], check=True)
    base = os.path.splitext(os.path.basename(doc_path))[0]
    return os.path.join(output_dir, f"{base}.docx")

def extract_text_from_docx(docx_path, output_dir):
    print("[+] Extracting plain text...")
    doc = Document(docx_path)
    with open(os.path.join(output_dir, "extracted_text.txt"), "w", encoding="utf-8") as f:
        for para in doc.paragraphs:
            f.write(para.text + "\n")

def extract_metadata(docx_path, output_dir):
    print("[+] Extracting metadata...")
    try:
        with zipfile.ZipFile(docx_path) as z:
            for name in ["docProps/core.xml", "docProps/app.xml"]:
                if name in z.namelist():
                    content = z.read(name).decode("utf-8")
                    with open(os.path.join(output_dir, f"meta_{os.path.basename(name)}.xml"), "w", encoding="utf-8") as f:
                        f.write(content)
    except Exception as e:
        print(f"[!] Metadata extraction failed: {e}")

def extract_urls(docx_path, output_dir):
    print("[+] Extracting hyperlinks...")
    urls = set()
    doc = Document(docx_path)
    for para in doc.paragraphs:
        found = re.findall(r'https?://[^\s\\\"]+', para.text)
        urls.update(found)
    with open(os.path.join(output_dir, "extracted_urls.txt"), "w") as f:
        for u in sorted(urls):
            f.write(u + "\n")

def extract_media(docx_path, output_dir, export_svg=True, export_png=True):
    with zipfile.ZipFile(docx_path, 'r') as z:
        for file in z.namelist():
            if file.startswith("word/media/") or file.startswith("word/embeddings/"):
                filename = os.path.basename(file)
                target_path = os.path.join(output_dir, filename)
                with open(target_path, "wb") as f:
                    f.write(z.read(file))
                if filename.lower().endswith(".wmf"):
                    if export_svg:
                        subprocess.run([LIBREOFFICE_PATH, "--headless", "--convert-to", "svg", "--outdir", output_dir, target_path])
                    if export_png:
                        subprocess.run([LIBREOFFICE_PATH, "--headless", "--convert-to", "png", "--outdir", output_dir, target_path])

def render_doc_pages(doc_path, output_dir):
    print("[+] Rendering full pages...")
    subprocess.run([LIBREOFFICE_PATH, "--headless", "--convert-to", "png", "--outdir", output_dir, doc_path])

def run_ocr(output_dir):
    print("[+] Running OCR...")
    for file in os.listdir(output_dir):
        if file.lower().endswith(".png"):
            path = os.path.join(output_dir, file)
            text = pytesseract.image_to_string(Image.open(path))
            with open(path + ".ocr.txt", "w", encoding="utf-8") as f:
                f.write(text)

def process_file(filepath, render_pages=True, do_ocr=False, export_svg=True, export_png=True):
    base = os.path.splitext(os.path.basename(filepath))[0]
    outdir = os.path.join(os.path.dirname(filepath), f"ripped_output_{base}")
    os.makedirs(outdir, exist_ok=True)

    ext = os.path.splitext(filepath)[1].lower()
    if ext == ".doc":
        filepath = convert_doc_to_docx(filepath, outdir)

    extract_text_from_docx(filepath, outdir)
    extract_metadata(filepath, outdir)
    extract_urls(filepath, outdir)
    extract_media(filepath, outdir, export_svg, export_png)
    if render_pages:
        render_doc_pages(filepath, outdir)
    if do_ocr:
        run_ocr(outdir)
    print(f"[✓] Done! Output in {outdir}")

# GUI & CLI handled in next part...
